"""Markdown to DOCX conversion using the Markdown token stream."""

from __future__ import annotations

from collections.abc import Sequence
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from .errors import ExportConversionError
from .markdown_ast import inline_plain_text, parse_markdown, token_attrs


FONT_NAME = "LXGW WenKai Lite"


def _safe_href(href: str) -> str | None:
    parsed = urlparse(href)
    if parsed.scheme.lower() in {"http", "https", "mailto"}:
        return href
    if not parsed.scheme and not href.startswith("//"):
        return href
    return None


def _set_run_font(run: Any, *, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = "Consolas" if code else FONT_NAME
    run.font.size = Pt(9.5 if code else 10.5)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    selected = "Consolas" if code else FONT_NAME
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attribute}"), selected)


def _add_hyperlink(paragraph: Any, text: str, href: str, *, bold: bool, italic: bool) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RELATIONSHIP
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    safe_href = _safe_href(href)
    if not safe_href:
        run = paragraph.add_run(text)
        _set_run_font(run, bold=bold, italic=italic)
        return
    relation_id = paragraph.part.relate_to(safe_href, RELATIONSHIP.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    if italic:
        run_properties.append(OxmlElement("w:i"))
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    run_properties.append(fonts)
    run_element.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _add_inline(paragraph: Any, token: Any, *, force_bold: bool = False) -> None:
    children = token.children or ()
    bold = force_bold
    italic = False
    code = False
    links: list[str] = []
    for child in children:
        kind = child.type
        if kind == "strong_open":
            bold = True
        elif kind == "strong_close":
            bold = force_bold
        elif kind in {"em_open", "s_open"}:
            italic = True
        elif kind in {"em_close", "s_close"}:
            italic = False
        elif kind == "code_inline":
            run = paragraph.add_run(child.content)
            _set_run_font(run, bold=bold, italic=italic, code=True)
        elif kind == "link_open":
            links.append(token_attrs(child).get("href", ""))
        elif kind == "link_close":
            if links:
                links.pop()
        elif kind in {"text", "html_inline", "image"}:
            text = child.content
            if kind == "image":
                text = token_attrs(child).get("alt", child.content)
            if not text:
                continue
            if links:
                _add_hyperlink(paragraph, text, links[-1], bold=bold, italic=italic)
            else:
                run = paragraph.add_run(text)
                _set_run_font(run, bold=bold, italic=italic, code=code)
        elif kind in {"softbreak", "hardbreak"}:
            run = paragraph.add_run("\n")
            _set_run_font(run, bold=bold, italic=italic)


def _style_font(document: Any) -> None:
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for style in document.styles:
        if not hasattr(style, "font"):
            continue
        style.font.name = FONT_NAME
        style.font.size = Pt(10.5)
        element = style._element
        rpr = element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), FONT_NAME)
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)


class _DocxRenderer:
    def __init__(self, document: Any) -> None:
        self.document = document
        self.list_stack: list[str] = []
        self.blockquote_depth = 0

    def _paragraph(self, *, style: str | None = None) -> Any:
        paragraph_style = style
        if paragraph_style is None and self.list_stack:
            depth = min(len(self.list_stack), 3)
            list_kind = self.list_stack[-1]
            base_style = "List Number" if list_kind == "ordered" else "List Bullet"
            paragraph_style = base_style if depth == 1 else f"{base_style} {depth}"
        try:
            paragraph = (
                self.document.add_paragraph(style=paragraph_style)
                if paragraph_style
                else self.document.add_paragraph()
            )
        except KeyError:
            paragraph = self.document.add_paragraph()
        if self.list_stack:
            from docx.shared import Inches

            paragraph.paragraph_format.left_indent = Inches(0.25 * len(self.list_stack))
        if self.blockquote_depth and not self.list_stack:
            try:
                paragraph.style = "Intense Quote"
            except KeyError:
                pass
        return paragraph

    def _render_table(self, tokens: Sequence[Any], start: int) -> int:
        rows: list[list[tuple[Any | None, bool]]] = []
        current_row: list[tuple[Any | None, bool]] | None = None
        current_cell: Any | None = None
        is_header = False
        index = start + 1
        while index < len(tokens):
            token = tokens[index]
            kind = token.type
            if kind == "tr_open":
                current_row = []
            elif kind in {"th_open", "td_open"}:
                current_cell = None
                is_header = kind == "th_open"
            elif kind == "inline" and current_row is not None:
                current_cell = token
            elif kind in {"th_close", "td_close"} and current_row is not None:
                current_row.append((current_cell, is_header))
                current_cell = None
            elif kind == "tr_close" and current_row is not None:
                rows.append(current_row)
                current_row = None
            elif kind == "table_close":
                break
            index += 1

        if rows:
            table = self.document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            for row_index, row in enumerate(rows):
                for col_index, (inline, header) in enumerate(row):
                    paragraph = table.cell(row_index, col_index).paragraphs[0]
                    if inline is not None:
                        _add_inline(paragraph, inline, force_bold=header)
        return index

    def render(self, tokens: Sequence[Any]) -> None:
        index = 0
        while index < len(tokens):
            token = tokens[index]
            kind = token.type
            if kind == "heading_open":
                level = max(1, min(6, int(token.tag[1:] or "1")))
                inline = tokens[index + 1] if index + 1 < len(tokens) else None
                paragraph = self._paragraph(style=f"Heading {level}")
                if inline is not None and inline.type == "inline":
                    _add_inline(paragraph, inline)
                index += 3
                continue
            if kind == "paragraph_open":
                inline = tokens[index + 1] if index + 1 < len(tokens) else None
                paragraph = self._paragraph()
                if inline is not None and inline.type == "inline":
                    _add_inline(paragraph, inline)
                index += 3
                continue
            if kind in {"bullet_list_open", "ordered_list_open"}:
                self.list_stack.append("ordered" if kind.startswith("ordered") else "bullet")
            elif kind in {"bullet_list_close", "ordered_list_close"}:
                if self.list_stack:
                    self.list_stack.pop()
            elif kind == "blockquote_open":
                self.blockquote_depth += 1
            elif kind == "blockquote_close":
                self.blockquote_depth = max(0, self.blockquote_depth - 1)
            elif kind in {"fence", "code_block"}:
                paragraph = self._paragraph()
                run = paragraph.add_run(token.content.rstrip("\n"))
                _set_run_font(run, code=True)
            elif kind == "table_open":
                index = self._render_table(tokens, index)
            elif kind == "hr":
                paragraph = self._paragraph()
                run = paragraph.add_run("――――")
                _set_run_font(run)
            index += 1


def markdown_to_docx(content: str) -> bytes:
    """Convert Markdown into a DOCX byte stream.

    The function is deliberately independent from the Tool and HTTP layers so
    it can be tested with a Markdown fixture and reused by other adapters.
    """

    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ExportConversionError(
            "docx_dependency_missing", "python-docx is required for DOCX export"
        ) from exc

    try:
        document = Document()
        for section in document.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)
        _style_font(document)
        _DocxRenderer(document).render(parse_markdown(content))
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except ExportConversionError:
        raise
    except Exception as exc:
        raise ExportConversionError("docx_conversion_failed", str(exc) or "DOCX conversion failed") from exc


__all__ = ["markdown_to_docx"]
