from __future__ import annotations

import asyncio
from pathlib import Path

import pytest

import server.services.exporters.pdf_exporter as pdf_exporter
from server.agent.memory import InMemoryConversationStore
from server.agent.registry import ToolExecutionContext
from server.agent.tools import export_file_tool
from server.services.document_exporter import (
    MAX_MARKDOWN_BYTES,
    DocumentExportError,
    DocumentExporter,
    sanitize_filename,
)
from server.services.exporters.markdown_ast import (
    markdown_to_plain_text,
    normalize_markdown_for_export,
    parse_markdown,
)
from server.services.exporters.errors import ExportConversionError
from server.services.exporters.pdf_exporter import _font_only_url_fetcher, markdown_to_html


MARKDOWN_FIXTURE = """# 导出报告

这是一个**粗体**和*斜体*段落，包含[链接](https://example.com)。

- 无序项目
- 第二项

1. 有序项目
2. 第二项

> 这是一段引用。

```python
print('你好')
```

| 名称 | 值 |
| --- | --- |
| 中文 | 42 |
"""

REGRESSION_FIXTURE = (
    Path(__file__).resolve().parents[1] / "fixtures" / "export_file" / "定语从句教案v1.md"
)


def test_markdown_to_plain_text_removes_markup_but_keeps_blocks() -> None:
    rendered = markdown_to_plain_text(MARKDOWN_FIXTURE)

    assert "# " not in rendered
    assert "**" not in rendered
    assert "导出报告" in rendered
    assert "- 无序项目" in rendered
    assert "1. 有序项目" in rendered
    assert "print('你好')" in rendered
    assert "中文 | 42" in rendered
    assert "https://example.com" in rendered


def test_list_tail_plain_text_is_parsed_as_a_new_paragraph() -> None:
    source = "- 环节一\n提问：这是什么？\n"

    assert normalize_markdown_for_export(source) == "- 环节一\n\n提问：这是什么？\n"

    token_types = [token.type for token in parse_markdown(source)]
    assert token_types == [
        "bullet_list_open",
        "list_item_open",
        "paragraph_open",
        "inline",
        "paragraph_close",
        "list_item_close",
        "bullet_list_close",
        "paragraph_open",
        "inline",
        "paragraph_close",
    ]
    html = markdown_to_html(source)
    assert "<li>环节一</li>\n</ul>\n<p>提问：这是什么？</p>" in html


@pytest.mark.parametrize(
    ("source",),
    [
        ("- 第一项\n- 第二项\n",),
        ("- 父项\n  - 子项\n    子项续行\n- 同级项\n",),
        ("- 项目\n  显式缩进的续行\n",),
        ("- 项目\n```text\n代码块内容\n```\n",),
        ("```text\n- 看起来像列表的代码\n普通代码文本\n```\n",),
        ("    - 看起来像列表的代码\n    普通代码文本\n",),
        ("- 项目\n# 标题\n",),
        ("- 项目\n> 引用\n",),
        ("- 项目\n| 名称 | 值 |\n| --- | --- |\n| A | B |\n",),
        ("- 项目\nSetext 标题\n---\n",),
    ],
)
def test_list_tail_normalization_leaves_nonparagraph_boundaries_untouched(source: str) -> None:
    assert normalize_markdown_for_export(source) == source


def test_real_agent_lesson_plan_regression_exports_list_tail_as_paragraph(tmp_path: Path) -> None:
    source = REGRESSION_FIXTURE.read_text(encoding="utf-8")
    list_item = "- The students were rescued by the rescuers. The rescuers arrived quickly."
    question = '提问："这两句话都在说同一批人，能不能合并成一句更简洁的表达？" 引出"修饰语"与"定语"概念。'

    # This fixture is the unmodified agent output that originally reproduced
    # the PDF defect: there is intentionally no blank line at this boundary.
    assert f"{list_item}\n{question}" in source
    normalized = normalize_markdown_for_export(source)
    assert f"{list_item}\n\n{question}" in normalized

    tokens = parse_markdown(source)
    list_inline_index = next(
        index
        for index, token in enumerate(tokens)
        if token.type == "inline" and token.content == list_item.removeprefix("- ")
    )
    assert tokens[list_inline_index + 3].type == "bullet_list_close"
    assert tokens[list_inline_index + 5].type == "inline"
    assert tokens[list_inline_index + 5].content == question

    html = markdown_to_html(source)
    assert f"<li>{list_item.removeprefix('- ')}</li>\n</ul>\n<p>提问：" in html

    exporter = DocumentExporter()
    outputs = {
        format_name: exporter.export(
            filename="lesson-plan",
            format=format_name,
            content=source,
            workspace_root=tmp_path,
        )
        for format_name in ("md", "txt", "docx", "pdf")
    }
    assert outputs["md"].path.read_text(encoding="utf-8") == source

    text_output = outputs["txt"].path.read_text(encoding="utf-8")
    assert f"{list_item}\n\n{question}" in text_output

    from docx import Document

    document = Document(outputs["docx"].path)
    list_paragraph_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == list_item.removeprefix("- ")
    )
    question_paragraph_index = next(
        index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == question
    )
    assert document.paragraphs[list_paragraph_index].style.name == "List Bullet"
    assert document.paragraphs[question_paragraph_index].style.name != "List Bullet"
    assert question_paragraph_index == list_paragraph_index + 1

    pdf = outputs["pdf"].path.read_bytes()
    assert pdf.startswith(b"%PDF")
    assert len(pdf) > 1_000


def test_exporter_writes_all_formats_and_docx_structure(tmp_path: Path) -> None:
    exporter = DocumentExporter()
    outputs = {}
    for format_name in ("md", "txt", "docx", "pdf"):
        outputs[format_name] = exporter.export(
            filename="report",
            format=format_name,
            content=MARKDOWN_FIXTURE,
            workspace_root=tmp_path,
        )

    assert outputs["md"].path.read_text(encoding="utf-8") == MARKDOWN_FIXTURE
    assert "**粗体**" not in outputs["txt"].path.read_text(encoding="utf-8")
    assert outputs["md"].mime_type.startswith("text/markdown")
    assert outputs["txt"].filename == "report.txt"
    assert outputs["docx"].path.read_bytes().startswith(b"PK")
    assert outputs["pdf"].path.read_bytes().startswith(b"%PDF")

    from docx import Document

    document = Document(outputs["docx"].path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "导出报告" in paragraph_text
    assert "无序项目" in paragraph_text
    assert "粗体" in paragraph_text
    assert any(
        paragraph.style.name == "List Bullet"
        for paragraph in document.paragraphs
    )
    assert any(
        paragraph.style.name == "List Number"
        for paragraph in document.paragraphs
    )
    assert document.tables[0].cell(1, 0).text == "中文"
    assert document.tables[0].cell(1, 1).text == "42"


def test_export_ids_are_unique_and_filename_cannot_escape_workspace(tmp_path: Path) -> None:
    exporter = DocumentExporter()
    first = exporter.export(
        filename="../outside/../report",
        format="md",
        content="first",
        workspace_root=tmp_path,
    )
    second = exporter.export(
        filename="../outside/../report",
        format="md",
        content="second",
        workspace_root=tmp_path,
    )

    export_dir = tmp_path / ".agent-exports"
    assert first.file_id != second.file_id
    assert first.path.parent == export_dir
    assert second.path.parent == export_dir
    assert first.filename == second.filename
    assert not (tmp_path.parent / "outside").exists()
    assert exporter.resolve(first.file_id, tmp_path) == first


def test_export_limits_and_structured_tool_errors(tmp_path: Path) -> None:
    exporter = DocumentExporter()
    with pytest.raises(DocumentExportError, match="exceeds"):
        exporter.export(
            filename="too-large",
            format="md",
            content="x" * (MAX_MARKDOWN_BYTES + 1),
            workspace_root=tmp_path,
        )

    context = ToolExecutionContext(
        session_id="session",
        store=InMemoryConversationStore(),
        workspace_root=tmp_path,
        request_id="request",
    )
    tool = export_file_tool(exporter)
    error = asyncio.run(
        tool.executor(
            {"filename": "", "format": "pdf", "content": "# title"},
            context,
        )
    )
    assert isinstance(error, dict)
    assert error["error"] is True
    assert error["code"] == "invalid_filename"


def test_export_file_description_requires_a_blank_list_paragraph_boundary() -> None:
    description = export_file_tool().description

    assert "valid Markdown" in description
    assert "blank line" in description


def test_sanitize_filename_removes_unsafe_segments() -> None:
    safe = sanitize_filename("..\\..\\CON?.pdf", "pdf")
    assert "/" not in safe and "\\" not in safe
    assert ".." not in safe
    assert safe != "CON"


def test_pdf_url_fetcher_only_allows_bundled_woff2_files(tmp_path: Path) -> None:
    class FakeURLFetcher:
        def __init__(self, **kwargs) -> None:
            self.options = kwargs

        def fetch(self, url: str, headers=None):
            return {"url": url}

    class FakeFatalURLFetchingError(RuntimeError):
        pass

    font_directory = tmp_path / "fonts"
    font_directory.mkdir()
    bundled_font = font_directory / "files" / "font.woff2"
    bundled_font.parent.mkdir()
    bundled_font.write_bytes(b"font")
    outside_font = tmp_path / "outside.woff2"
    outside_font.write_bytes(b"outside")

    fetcher = _font_only_url_fetcher(
        font_directory,
        FakeURLFetcher,
        FakeFatalURLFetchingError,
    )

    assert fetcher.fetch(bundled_font.as_uri())["url"] == bundled_font.as_uri()
    with pytest.raises(FakeFatalURLFetchingError):
        fetcher.fetch(outside_font.as_uri())
    with pytest.raises(FakeFatalURLFetchingError):
        fetcher.fetch("https://example.com/font.woff2")


def test_pdf_font_assets_use_one_complete_woff2_without_unicode_ranges() -> None:
    font_directory = pdf_exporter._font_directory()
    css_path, font_path = pdf_exporter._pdf_font_assets(font_directory)

    css = css_path.read_text(encoding="utf-8")
    assert css_path.name == pdf_exporter.PDF_FONT_CSS_FILENAME
    assert font_path.name == "lxgwwenkailite-regular-full.woff2"
    assert font_path.name in css
    assert css.count("@font-face") == 1
    assert "unicode-range" not in css
    assert font_path.stat().st_size > 3_000_000

    browser_css = (font_directory / "lxgwwenkailite-regular.css").read_text(encoding="utf-8")
    assert "unicode-range" in browser_css


@pytest.mark.parametrize(
    "missing_path",
    [pdf_exporter.PDF_FONT_CSS_FILENAME, pdf_exporter.PDF_FONT_FILENAME],
)
def test_pdf_export_reports_missing_pdf_font_assets_stably(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, missing_path: str
) -> None:
    font_directory = tmp_path / "fonts"
    font_directory.mkdir()
    css_path = font_directory / pdf_exporter.PDF_FONT_CSS_FILENAME
    font_path = font_directory / pdf_exporter.PDF_FONT_FILENAME

    if missing_path != pdf_exporter.PDF_FONT_CSS_FILENAME:
        css_path.write_text("@font-face {}", encoding="utf-8")
    if missing_path != pdf_exporter.PDF_FONT_FILENAME:
        font_path.parent.mkdir()
        font_path.write_bytes(b"font")
    monkeypatch.setattr(pdf_exporter, "_font_directory", lambda: font_directory)

    with pytest.raises(ExportConversionError) as caught:
        pdf_exporter.markdown_to_pdf("# 中文")

    assert caught.value.code == "pdf_font_missing"
